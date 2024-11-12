from pytubefix import YouTube
import whisper
from datetime import timedelta
import google.generativeai as genai
import os
import re
import json
from docx import Document

def download_audio(url, filename="Audio.mp3"):
    yt = YouTube(url)
    yt.streams.get_audio_only().download(filename=filename)
    print(f"Audio downloaded and saved as {filename}")
    return filename  

def format_time(seconds):
    return str(timedelta(seconds=round(seconds, 2)))

def transcribe_audio(path_to_audio="Audio.mp3"):
    model = whisper.load_model("tiny")
    result = model.transcribe(path_to_audio)
    data = ""
    for segment in result["segments"]:
        text = segment["text"]
        start_time = format_time(segment['start'])
        end_time = format_time(segment['end'])
        data += f'{text} Start Time: {start_time} End Time: {end_time}\n'
    print("Transcription complete.")
    return data  

def generate_json_from_transcript(transcript_data):
    prompt = (
    "You are an expert in creating structured JSON responses from video transcripts to be converted into detailed Word documents. "
    "Format the JSON to include a 'title' field, followed by a 'table_of_contents' field as an array listing the main sections. "
    "Then use a 'sections' field as an array, where each object includes:\n"
    "1. A 'heading' field for the main section title,\n"
    "2. A 'content' field with relevant quotes and timestamps (e.g., [00:05:12]) based on the transcript content. Skip any sponsored segments and focus only on the main content of the video.\n\n"
    "Format Example:\n"
    "{\n"
    "    'title': 'Document Title',\n"
    "    'table_of_contents': ['Section 1', 'Section 2'],\n"
    "    'sections': [\n"
    "        {\n"
    "            'heading': 'Section 1',\n"
    "            'content': '...'\n"
    "        },\n"
    "        {\n"
    "            'heading': 'Section 2',\n"
    "            'content': '...'\n"
    "        }\n"
    "    ]\n"
    "}\n\n"
    f"Video Transcript: {transcript_data}\n"
    "JSON Output:"
    )
    genai.configure(api_key=os.environ["GEMINI_API_KEY"])
    model = genai.GenerativeModel(model_name="gemini-1.5-pro-002")
    response = model.generate_content(prompt)
    cleaned_content = re.sub(r"```json\s*|\s*```", "", response.text, flags=re.DOTALL).strip()
    print("JSON generation from transcript complete.")
    return cleaned_content

def generate_word_document(json_data):
    data = json.loads(json_data)
    doc = Document()
    doc.add_heading(data['title'], level=1)
    doc.add_heading("Table of Contents", level=2)
    for i, toc_item in enumerate(data['table_of_contents'], start=1):
        doc.add_paragraph(f"{i}. {toc_item}")
    for section in data['sections']:
        doc.add_heading(section['heading'], level=2) 
        doc.add_paragraph(section['content'])
    doc_title = data['title'].replace(" ", "_") + ".docx"
    doc.save(doc_title)
    print(f"Document '{doc_title}' created successfully!")
    return doc_title  

def process_youtube_to_document(yt_url):
    audio_file = download_audio(yt_url)
    transcript = transcribe_audio(audio_file)
    json_content = generate_json_from_transcript(transcript)
    document_name = generate_word_document(json_content)
    return document_name 


yt_url = "https://www.youtube.com/watch?v=ktkjUjcZid0"  
final_document = process_youtube_to_document(yt_url)
print(f"Process completed. Final document saved as: {final_document}")
