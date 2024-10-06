import os
import gradio as gr
import time

from langchain_community.llms import ollama
from langchain.prompts import PromptTemplate
from pytubefix import YouTube
import whisper
import markdown
from docx import Document
from bs4 import BeautifulSoup

def download_video_audio(link):
    """
    Downloads the audio from YouTube using pytubefix.
    """
    yt = YouTube(link)
    ys = yt.streams.get_audio_only()
    ys.download(filename="Audio", mp3=True)

def generate_transcription():
    """
    Uses Whisper model large-v3-turbo to generate the transcription with timestamps.
    """
    model = whisper.load_model("large-v3-turbo")
    result = model.transcribe("Audio.mp3")
    transcription = ""
    for segment in result["segments"]:
        transcription += f"{segment['text']} (Start: {round(segment['start'],2)}s - End: {round(segment['end'],2)}s)\n"
    return transcription

def doc_generation(data):
    ollama_base_url = os.getenv("OLLAMA_BASE_URL")
    model = ollama.Ollama(
        base_url=ollama_base_url,
        model='llama3.1',
    )
    prompt_template = PromptTemplate.from_template(
        """Here is a video Transcript {data}. Based on this, provide a detailed summary document of the video with in-depth paragraphs and timestamps that can be easily referred to by the user.
Generated Document: """
    )
    prompt = prompt_template.format(data=data)
    completion = model.invoke(prompt, temperature=0)
    return completion

def markdown_to_docx(markdown_text, output_file):
    """
    Uses Ollama to generate a markdown
    """
    html = markdown.markdown(markdown_text)
    soup = BeautifulSoup(html, "html.parser")
    doc = Document()
    for element in soup:
        if element.name == 'p':
            paragraph = doc.add_paragraph()
            for part in element:
                if part.name == 'strong':
                    run = paragraph.add_run(part.get_text())
                    run.bold = True
                else:
                    paragraph.add_run(part)
    doc.save(output_file)

def process_link(link):
    """
    Final Function that call the other functions
    """
    download_video_audio(link)
    transcription = generate_transcription()
    markdown_text = doc_generation(transcription)
    output_file = 'summary_document.docx'
    try:
        markdown_to_docx(markdown_text, output_file)
        result_message = f"Word document '{output_file}' created successfully."
    except Exception as e:
        result_message = f"Caught an exception: {e}"
    return output_file


def gradio_interface(theme=gr.themes.Soft(), css="footer {display: none !important;} "):
    with gr.Blocks() as demo:
        gr.Markdown("""
        # 🎥 YouTube Video Summarizer 🎥
        Enter a YouTube link below to generate a detailed summary document with timestamps. The generated document is ready for easy reference!
        """)
        link_input = gr.Textbox(label="YouTube Link", placeholder="Enter YouTube link here...")
        file_output = gr.File(label="Download Summary Document", interactive=False)

        submit_button = gr.Button("Generate Summary")
        submit_button.click(fn=process_link, inputs=[link_input], outputs=[file_output])
    return demo

if __name__ == "__main__":
    demo = gradio_interface()
    demo.launch()